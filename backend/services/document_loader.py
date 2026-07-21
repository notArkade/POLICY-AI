# from pathlib import Path

# from docx import Document
# from langchain_community.document_loaders import PyPDFLoader


# def _validate_file(file_path, expected_suffix):
#     path = Path(file_path)

#     if not path.exists():
#         raise FileNotFoundError(f"File not found: {path}")

#     if not path.is_file():
#         raise ValueError(f"Expected a file path, got: {path}")

#     if path.suffix.lower() != expected_suffix:
#         raise ValueError(f"Expected a {expected_suffix} file, got: {path.suffix}")

#     return path


# def load_pdf(file_path):
#     path = _validate_file(file_path, ".pdf")
#     loader = PyPDFLoader(str(path))
#     docs = loader.load()

#     text = "\n".join(doc.page_content.strip() for doc in docs if doc.page_content)
#     if not text:
#         raise ValueError(f"No extractable text found in PDF: {path}")

#     return text


# def load_docx(file_path):
#     path = _validate_file(file_path, ".docx")
#     doc = Document(str(path))

#     text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
#     if not text:
#         raise ValueError(f"No extractable text found in DOCX: {path}")

#     return text

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from langchain_community.document_loaders import PyPDFLoader
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".doc"}


def _validate_file(file_path, expected_suffix=None):
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Expected a file path, got: {path}")

    suffix = path.suffix.lower()
    if expected_suffix and suffix != expected_suffix:
        raise ValueError(f"Expected a {expected_suffix} file, got: {path.suffix}")

    if not expected_suffix and suffix not in SUPPORTED_SUFFIXES:
        allowed = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ValueError(f"Unsupported document type {path.suffix}. Allowed: {allowed}")

    return path


def _ensure_text(text, path):
    normalized = text.strip()
    if not normalized:
        raise ValueError(f"No extractable text found in document: {path}")
    return normalized


def load_pdf(file_path):
    path = _validate_file(file_path, ".pdf")
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    text = "\n".join(doc.page_content.strip() for doc in docs if doc.page_content)
    return _ensure_text(text, path)


def load_docx(file_path):
    path = _validate_file(file_path, ".docx")
    doc = Document(str(path))
    text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
    return _ensure_text(text, path)


def load_doc(file_path):
    path = _validate_file(file_path, ".doc")
    data = path.read_bytes()
    candidates = []

    for encoding in ("utf-16le", "latin-1"):
        decoded = data.decode(encoding, errors="ignore")
        matches = re.findall(r"[A-Za-z0-9][\w\s.,;:!?()/%&\-]{20,}", decoded)
        candidates.extend(match.strip() for match in matches)

    text = "\n".join(dict.fromkeys(candidates))
    return _ensure_text(text, path)


def load_document(file_path):
    path = _validate_file(file_path)

    if path.suffix.lower() == ".pdf":
        return load_pdf(path)
    if path.suffix.lower() == ".docx":
        return load_docx(path)
    if path.suffix.lower() == ".doc":
        return load_doc(path)

    raise ValueError(f"Unsupported document type: {path.suffix}")


def load_document_bytes(file_bytes: bytes, file_name: str):
    """Extract text from an upload held entirely in memory.

    Original files are persisted in Supabase Storage; no server-local copy is
    created merely to index a document.
    """
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        allowed = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ValueError(f"Unsupported document type {suffix}. Allowed: {allowed}")

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages)
    elif suffix == ".docx":
        document = Document(BytesIO(file_bytes))
        text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    else:
        candidates = []
        for encoding in ("utf-16le", "latin-1"):
            decoded = file_bytes.decode(encoding, errors="ignore")
            candidates.extend(re.findall(r"[A-Za-z0-9][\w\s.,;:!?()/%&\-]{20,}", decoded))
        text = "\n".join(dict.fromkeys(candidate.strip() for candidate in candidates))

    return _ensure_text(text, file_name)


def get_page_texts_from_bytes(file_bytes: bytes, file_name: str):
    """Return PDF page text for chunk citation metadata (one-based pages)."""
    if Path(file_name).suffix.lower() != ".pdf":
        return []
    return [(page.extract_text() or "").strip() for page in PdfReader(BytesIO(file_bytes)).pages]


def load_all_documents():
    """
    Loads all PDF and DOCX files from backend/documents
    Returns:
        [
            {
                "file_name": "...",
                "file_path": "...",
                "content": "..."
            }
        ]
    """

    docs_dir = Path(__file__).parent.parent / "documents"

    if not docs_dir.exists():
        raise FileNotFoundError(
            f"Documents folder not found: {docs_dir}"
        )

    documents = []

    for file in docs_dir.iterdir():

        try:
            if file.suffix.lower() == ".docx":

                documents.append(
                    {
                        "file_name": file.name,
                        "file_path": str(file),
                        "content": load_docx(file)
                    }
                )

            elif file.suffix.lower() == ".pdf":

                documents.append(
                    {
                        "file_name": file.name,
                        "file_path": str(file),
                        "content": load_pdf(file)
                    }
                )

        except Exception as e:
            print(f"Error loading {file.name}: {e}")

    return documents
